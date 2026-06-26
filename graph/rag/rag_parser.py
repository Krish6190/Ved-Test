import os
import zipfile
import tempfile
from .rag_splitter import RagTextSplitter

class RagDocumentParser:
    def __init__(self):
        self.text_splitter = RagTextSplitter()

    def process_file_to_chunks(self, file_path: str) -> list:
        """Reads a target file and converts it into a list of text chunks.

        Routes by extension:
          .pdf           -> pypdf text extraction
          .docx / .doc   -> python-docx paragraph extraction
          .zip           -> recursive ingest of contained files
          anything else  -> text fallback (utf-8 with errors=ignore)
        """
        if not os.path.exists(file_path):
            print(f"[RAG Error] Parser target path missing: {file_path}")
            return []

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == ".pdf":
                return self._read_pdf(file_path)
            if ext == ".docx":
                return self._read_docx(file_path)
            if ext == ".doc":
                # Legacy .doc: try python-docx (it handles some), otherwise fall back to text.
                try:
                    return self._read_docx(file_path)
                except Exception:
                    return self._read_text(file_path)
            if ext == ".zip":
                return self._read_zip(file_path)
            return self._read_text(file_path)
        except Exception as e:
            print(f"[RAG Error] {ext} parse failure for {file_path}: {e}")
            return []

    # ---------- readers ----------
    def _read_text(self, file_path: str) -> list:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        return self.text_splitter.split_raw_text(raw)

    def _read_pdf(self, file_path: str) -> list:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        parts = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt.strip():
                parts.append(txt)
        return self.text_splitter.split_raw_text("\n\n".join(parts))

    def _read_docx(self, file_path: str) -> list:
        import docx
        doc = docx.Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # Also pull text from tables (often holds real content in docs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text)
        return self.text_splitter.split_raw_text("\n\n".join(parts))

    def _read_zip(self, file_path: str) -> list:
        all_chunks = []
        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                with tempfile.TemporaryDirectory() as tmp:
                    for name in zf.namelist():
                        # Skip directories, hidden files, and known junk
                        if name.endswith("/"):
                            continue
                        if os.path.basename(name).startswith("."):
                            continue
                        if "__MACOSX" in name:
                            continue
                        try:
                            zf.extract(name, tmp)
                            extracted = os.path.join(tmp, name)
                            all_chunks.extend(self.process_file_to_chunks(extracted))
                        except Exception as exc:
                            print(f"[RAG Warning] Skipped {name}: {exc}")
                            continue
        except Exception as e:
            print(f"[RAG Error] ZIP read failure for {file_path}: {e}")
        return all_chunks
